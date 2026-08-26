"""Builds the site-progress Monthly Report .docx.

Python port of the `docx` (JS) rendering section of the original
site-report-generator.js script — same layout, fonts and colors (matching
the Khales "Monthly Report" Word template), rebuilt with python-docx so it
can run inside Odoo instead of Node.
"""

import io

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Pt, RGBColor
from PIL import Image

FONT_SERIF = "Times New Roman"
FONT_SANS = "Calibri"
COLOR_TEXT = "252525"
COLOR_GRAY = "808080"
COLOR_BORDER = "D9D9D9"
COLOR_ACCENT = "AF8D56"

EMU_PER_PIXEL = 9525  # 96 DPI, matches docx.js's pixel-based transformation


def _hp(half_points):
    """OOXML run size is in half-points."""
    return Pt(half_points / 2)


def _dxa(dxa):
    """OOXML spacing is in twentieths of a point (dxa)."""
    return Pt(dxa / 20)


def _color(hex_str):
    return RGBColor.from_string(hex_str)


def _set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_paragraph_bottom_border(paragraph, hex_color, size=6, space=4):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_paragraph_top_border(paragraph, hex_color, size=4, space=4):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), str(size))
    top.set(qn("w:space"), str(space))
    top.set(qn("w:color"), hex_color)
    pBdr.append(top)
    pPr.append(pBdr)


def _add_field(paragraph, field_code):
    """Insert a Word field (e.g. PAGE / NUMPAGES) as a run."""
    run = paragraph.add_run()
    r = run._r
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_end)
    return run


def prepare_embedded_photo(raw_bytes, max_width=300, max_height=220):
    """Downscale a photo to a sane embed size (JPEG q=80, capped 900x900),
    returning the box (in pixels, aspect preserved) to render it at.
    """
    with Image.open(io.BytesIO(raw_bytes)) as img:
        img = img.convert("RGB")
        img.thumbnail((900, 900))
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=80)
        width, height = img.size

    scale = min(max_width / width, max_height / height, 1)
    return {
        "data": buf.getvalue(),
        "width": round(width * scale),
        "height": round(height * scale),
    }


def _add_heading(document, text, with_rule=False):
    p = document.add_paragraph()
    p.paragraph_format.space_before = _dxa(300)
    p.paragraph_format.space_after = _dxa(120)
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_SANS
    run.font.size = _hp(24)
    run.font.color.rgb = _color(COLOR_TEXT)
    if with_rule:
        _set_paragraph_bottom_border(p, COLOR_ACCENT, size=6, space=4)
    return p


def _add_body_paragraph(document, text, gray=False, italic=False):
    p = document.add_paragraph()
    p.paragraph_format.space_after = _dxa(120)
    run = p.add_run(text or "")
    run.font.name = FONT_SANS
    run.font.size = _hp(20)
    run.font.color.rgb = _color(COLOR_GRAY if gray else COLOR_TEXT)
    run.italic = italic
    return p


def _add_bullet(document, text):
    p = document.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = _dxa(80)
    run = p.add_run(text)
    run.font.name = FONT_SANS
    run.font.size = _hp(20)
    run.font.color.rgb = _color(COLOR_TEXT)
    return p


def _add_info_table(document, rows):
    table = document.add_table(rows=0, cols=2)
    table.autofit = True
    for label, value in rows:
        row = table.add_row()
        label_cell, value_cell = row.cells
        _set_cell_shading(label_cell, COLOR_BORDER)
        label_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = label_cell.paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.name = FONT_SANS
        run.font.size = _hp(20)
        run.font.color.rgb = _color(COLOR_TEXT)

        value_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = value_cell.paragraphs[0]
        run = p.add_run(value or "")
        run.font.name = FONT_SANS
        run.font.size = _hp(20)
        run.font.color.rgb = _color(COLOR_TEXT)
    return table


def _add_photo_grid(document, photos):
    """2-column photo grid for one site visit. No per-photo captions — the
    narrative text already covers what happened at that visit.
    """
    prepared = [prepare_embedded_photo(p) for p in photos]
    rows = [prepared[i : i + 2] for i in range(0, len(prepared), 2)]

    table = document.add_table(rows=0, cols=2)
    for row_photos in rows:
        row = table.add_row()
        for i, cell in enumerate(row.cells):
            if i >= len(row_photos):
                continue
            photo = row_photos[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(
                io.BytesIO(photo["data"]),
                width=Emu(photo["width"] * EMU_PER_PIXEL),
                height=Emu(photo["height"] * EMU_PER_PIXEL),
            )
    return table


def _add_date_line(document, label):
    p = document.add_paragraph()
    p.paragraph_format.space_before = _dxa(100)
    p.paragraph_format.space_after = _dxa(200)
    run = p.add_run(f"Date: {label}")
    run.italic = True
    run.font.name = FONT_SANS
    run.font.size = _hp(18)
    run.font.color.rgb = _color(COLOR_GRAY)
    return p


def _add_footer(document, logo_path=None):
    section = document.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.paragraph_format.tab_stops.add_tab_stop(Pt(468), WD_TAB_ALIGNMENT.RIGHT)
    _set_paragraph_top_border(p, COLOR_BORDER, size=4, space=4)

    def run_of(text=""):
        r = p.add_run(text)
        r.font.name = FONT_SANS
        r.font.size = _hp(14)
        r.font.color.rgb = _color(COLOR_GRAY)
        return r

    run_of("Confidential - Prepared for Owner Use Only")
    run_of("\tPage ")
    _add_field(p, "PAGE")
    run_of(" of ")
    _add_field(p, "NUMPAGES")


def _add_cover_page(document, project, period_label, logo_path):
    p = document.add_paragraph()
    run = p.add_run(period_label)
    run.font.name = FONT_SERIF
    run.font.size = _hp(24)
    run.font.color.rgb = _color(COLOR_TEXT)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = _dxa(1600)

    logo_p = document.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if logo_path:
        logo_p.add_run().add_picture(logo_path, width=Emu(150 * EMU_PER_PIXEL), height=Emu(71 * EMU_PER_PIXEL))

    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = _dxa(500)
    run = title_p.add_run("Monthly Report")
    run.font.name = FONT_SERIF
    run.font.size = _hp(72)

    sub_p = document.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_before = _dxa(150)
    run = sub_p.add_run("Khales Engineering Consultancy")
    run.font.name = FONT_SERIF
    run.font.size = _hp(24)

    gap = document.add_paragraph()
    gap.paragraph_format.space_before = _dxa(3200)

    for label, value in [
        ("Prepared For", f"Mr. {project.get('client_name') or ''}"),
        ("Project Location", project.get("location") or ""),
        ("Plot Number", project.get("plot_number") or ""),
        ("Document Ref", project.get("project_no") or ""),
    ]:
        line = document.add_paragraph()
        run = line.add_run(f"{label}: {value}")
        run.font.name = FONT_SERIF
        run.font.size = _hp(20)
        run.font.color.rgb = _color(COLOR_TEXT)

    document.add_page_break()


def build_report_docx(project, period_label, visit_dates_label, visits, synthesis, logo_path=None):
    """project: dict with project_no, project_name, location, contractor,
    consultant, client_name, plot_number, manager_name.
    visits: list of {"date_label": str, "narrative": str, "photos": [bytes, ...]}.
    synthesis: dict from claude_synthesis.synthesize_monthly_report.
    Returns the .docx file content as bytes.
    """
    document = Document()
    _add_footer(document, logo_path)
    _add_cover_page(document, project, period_label, logo_path)

    _add_info_table(
        document,
        [
            ("Project Name", project.get("project_name")),
            ("Location", project.get("location") or ""),
            ("Contractor", project.get("contractor") or ""),
            ("Consultant", project.get("consultant") or ""),
        ],
    )
    _add_body_paragraph(
        document,
        (
            "This report was generated automatically: site-visit photographs and their AI-written summaries "
            "were pulled from the project's records, and this month's planned activities/recommendations "
            "were drafted by AI from those summaries. Please verify before formal client issue."
        ),
        gray=True,
        italic=True,
    )

    _add_heading(document, "1. EXECUTIVE SUMMARY")
    _add_body_paragraph(
        document,
        f"This report covers {len(visits)} site visit(s) for {project.get('project_name')} during "
        f"{period_label} ({visit_dates_label}).",
    )

    _add_heading(document, "2. SITE UPDATE - WORKS COMPLETED THIS MONTH", with_rule=True)
    for para in synthesis["site_update_summary"].split("\n\n"):
        if para.strip():
            _add_body_paragraph(document, para.strip())

    _add_heading(document, "3. SITE PHOTOS")
    for visit in visits:
        if visit["photos"]:
            _add_photo_grid(document, visit["photos"])
        _add_date_line(document, visit["date_label"])

    _add_heading(document, "4. PLANNED ACTIVITIES — NEXT MONTH", with_rule=True)
    for item in synthesis["planned_activities"]:
        _add_bullet(document, item)

    _add_heading(document, "5. RECOMMENDATIONS / OWNER ACTION REQUIRED")
    _add_body_paragraph(document, synthesis["recommendations"])

    sign_table = document.add_table(rows=1, cols=2)
    sign_table.rows[0].cells[0].paragraphs[0].add_run("Prepared By").bold = True
    p = sign_table.rows[0].cells[0].add_paragraph()
    run = p.add_run(f"{project.get('manager_name')} / Project Manager" if project.get("manager_name") else "[Name / Title / Signature]")
    run.italic = True
    run.font.color.rgb = _color(COLOR_GRAY)

    sign_table.rows[0].cells[1].paragraphs[0].add_run("Reviewed / Approved By").bold = True
    p = sign_table.rows[0].cells[1].add_paragraph()
    run = p.add_run("[Name / Title / Signature]")
    run.italic = True
    run.font.color.rgb = _color(COLOR_GRAY)

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
